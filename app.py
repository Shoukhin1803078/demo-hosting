from flask import Flask, request, jsonify, send_from_directory, send_file, url_for, make_response
import os
import uuid
import io
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import (
    ChatPromptTemplate,
    MessagesPlaceholder,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from docx import Document
from flask_caching import Cache
from werkzeug.exceptions import BadRequest, NotFound, InternalServerError
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

app = Flask(__name__)

# Set up caching
app.config['CACHE_TYPE'] = 'simple'
cache = Cache(app)

# Set up rate limiting
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://",
)


# Set the API key (use environment variables in production)
os.environ['OPENAI_API_KEY'] ="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"


SYSTEM_MESSAGE = """
You are "KUROCO LAB chatbot", created by JB Connect Ltd. As a managing director and project implementor, your role is to:
1. Guide users through project descriptions, asking relevant questions to gather comprehensive information.
2. Produce detailed project summaries covering: project name, goals, scope, timeline, budget, resources needed, risks, and key stakeholders.
3. Offer industry-specific insights and best practices for project management.
4. Provide approximate cost and timeline estimates based on project complexity.
5. Suggest potential challenges and mitigation strategies.
6. Recommend project management methodologies suitable for the described project.
7. Offer to create and provide downloadable SRS documents when appropriate.
Always maintain a professional yet approachable tone. Be proactive in seeking clarification and offering additional information to ensure comprehensive project planning.
"""

prompt = ChatPromptTemplate.from_messages([
    SystemMessagePromptTemplate.from_template(SYSTEM_MESSAGE),
    MessagesPlaceholder(variable_name="history"),
    HumanMessagePromptTemplate.from_template("{input}")
])

llm = ChatOpenAI(temperature=0.7)
memory = ConversationBufferMemory(return_messages=True)
conversation = ConversationChain(memory=memory, prompt=prompt, llm=llm)

documents = {}

def process_assistant_message(assistant_message, user_message):
    if any(keyword in user_message.lower() for keyword in ["document", "report", "summary", "download", "link", "srs"]):
        doc_id = str(uuid.uuid4())
        srs_content = generate_srs_content(memory.chat_memory.messages)
        documents[doc_id] = srs_content
        download_link = url_for('get_document', doc_id=doc_id, _external=True)
        assistant_message += f"\n\nI've prepared an SRS document based on our conversation. Here's the link to download your SRS document: <a href='{download_link}' target='_blank'>Download SRS Document</a>"
    return assistant_message

def generate_srs_content(messages):
    conversation_history = "\n".join([f"{msg.type}: {msg.content}" for msg in messages])
    
    srs_prompt = f"""
    Based on the following conversation, generate a comprehensive Software Requirements Specification (SRS) document. The structure and content should be entirely based on the information discussed in the conversation. Follow these guidelines:

    1. Start with an introduction that summarizes the project.
    2. Create logical sections based on the topics discussed in the conversation.
    3. Include all relevant details mentioned, such as project goals, scope, features, requirements, constraints, and any other important aspects.
    4. Use appropriate headings and subheadings to organize the information.
    5. If certain standard SRS sections are applicable but not explicitly discussed, include them with a note that they require further discussion.
    6. Ensure the document flows logically and covers all aspects of the project mentioned in the conversation.

    Conversation History:
    {conversation_history}

    Generate the SRS document content:
    """

    srs_response = llm.invoke(srs_prompt)
    return srs_response.content if hasattr(srs_response, 'content') else str(srs_response)

def create_srs_document(content):
    doc = Document()
    doc.add_heading('Software Requirements Specification (SRS)', 0)

    lines = content.split('\n')
    current_level = 0
    for line in lines:
        if line.strip():
            if line[0].isdigit() or line.isupper():
                level = len(line.split('.')) if '.' in line else (1 if line.isupper() else 2)
                doc.add_heading(line.strip(), level=level)
                current_level = level
            else:
                if line.startswith('  '):
                    doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line.strip())

    return doc

@app.route('/')
def home():
    return send_from_directory('.', 'index.html')

@app.route('/chat', methods=['POST'])
@limiter.limit("5 per minute")
def chat():
    try:
        user_message = request.json['message']
        if not user_message or not isinstance(user_message, str):
            raise BadRequest("Invalid message format")
        
        response = conversation.invoke(input=user_message)
        response_content = response['response'] if isinstance(response, dict) else str(response)
        processed_response = process_assistant_message(response_content, user_message)
        return jsonify({'response': processed_response})
    except Exception as e:
        app.logger.error(f"An error occurred: {str(e)}")
        raise InternalServerError("An unexpected error occurred")

@app.route("/create_document/<doc_id>", methods=["GET"])
def get_document(doc_id):
    try:
        if doc_id not in documents:
            raise NotFound("Document not found")
        content = documents[doc_id]
        doc = create_srs_document(content)
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='SRS_Document.docx'
        )
    except Exception as e:
        app.logger.error(f"An error occurred while creating the document: {e}")
        raise InternalServerError("Failed to create document")

@app.route('/export-chat', methods=['POST'])
def export_chat():
    try:
        chat_content = request.json['content']
        content_bytes = chat_content.encode('utf-8')
        
        # Check size (limit to 10 MB)
        max_size = 10 * 1024 * 1024  # 10 MB in bytes
        if len(content_bytes) > max_size:
            return jsonify({
                'error': f"Chat export is too large ({len(content_bytes) / 1024 / 1024:.2f} MB). Maximum size is {max_size / 1024 / 1024} MB."
            }), 413  # 413 Payload Too Large
        
        response = make_response(chat_content)
        response.headers.set('Content-Type', 'text/html')
        response.headers.set('Content-Disposition', 'attachment', filename='chat_export.html')
        return response

    except Exception as e:
        app.logger.error(f"An error occurred during chat export: {str(e)}")
        raise InternalServerError("An unexpected error occurred during chat export")

@app.errorhandler(BadRequest)
@app.errorhandler(NotFound)
@app.errorhandler(InternalServerError)
def handle_error(error):
    return jsonify({'error': str(error)}), error.code

if __name__ == '__main__':
    app.run(debug=True)